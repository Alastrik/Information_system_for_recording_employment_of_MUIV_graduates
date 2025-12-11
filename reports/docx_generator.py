# reports/docx_generator.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from database import DatabaseConnection
from datetime import datetime

class DocxReportGenerator:
    @staticmethod
    def generate_employment_report(output_path):
        """Генерирует .docx отчёт на основе данных из БД."""
        db = DatabaseConnection()
        db.connect()

        try:
            # Запрос: все выпускники и их трудоустройство
            query = """
                SELECT 
                    g.full_name,
                    g.graduation_year,
                    f.name AS faculty,
                    c.name AS company,
                    p.title AS position,
                    es.status_name AS status,
                    e.salary,
                    e.start_date,
                    e.is_current
                FROM graduate g
                LEFT JOIN user_profile up ON g.user_profile_id = up.id
                LEFT JOIN faculty f ON up.faculty_id = f.id
                LEFT JOIN employment e ON g.id = e.graduate_id
                LEFT JOIN company c ON e.company_id = c.id
                LEFT JOIN position p ON e.position_id = p.id
                LEFT JOIN employment_status es ON e.status_id = es.id
                ORDER BY g.full_name, e.start_date DESC
            """
            data = db.execute_query(query, fetch=True)

            # Создание документа
            doc = Document()
            doc.add_heading('Отчёт о трудоустройстве выпускников МУИВ', 0)
            doc.add_paragraph(f'Сформировано: {datetime.now().strftime("%d.%m.%Y %H:%M")}', style='Intense Quote')
            doc.add_paragraph()

            if not data:
                doc.add_paragraph("Нет данных для отчёта.")
            else:
                # Таблица
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                headers = ["ФИО", "Год\nвыпуска", "Факультет", "Организация", "Должность", "Статус", "Зарплата", "Текущее"]
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

                for row in data:
                    cells = table.add_row().cells
                    cells[0].text = row[0] or "—"
                    cells[1].text = str(row[1]) if row[1] else "—"
                    cells[2].text = row[2] or "—"
                    cells[3].text = row[3] or "—"
                    cells[4].text = row[4] or "—"
                    cells[5].text = row[5] or "—"
                    cells[6].text = f"{row[6]:,.2f} руб." if row[6] else "—"
                    cells[7].text = "Да" if row[8] else "Нет"

                # Автоподбор (ограничено в python-docx, но можно уменьшить шрифт)
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(8)

            # Подпись
            doc.add_paragraph()
            doc.add_paragraph("Автор: Монахов Артем Сергеевич", style='Intense Quote')

            # Сохранение
            import os
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return output_path

        finally:
            db.disconnect()